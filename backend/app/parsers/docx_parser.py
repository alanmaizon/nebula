from __future__ import annotations

import io
from pathlib import Path

from app.parsers.base import ParseResult, ParsedPage


class DocxDocumentParser:
    parser_id = "docx"
    _CONTENT_TYPES = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }

    def supports(self, *, file_name: str, content_type: str) -> bool:
        if content_type.lower() in self._CONTENT_TYPES:
            return True
        return Path(file_name).suffix.lower() == ".docx"

    def parse(self, *, content: bytes, file_name: str, content_type: str) -> ParseResult:
        del file_name, content_type
        try:
            from docx import Document
        except ImportError:
            return ParseResult(
                parser_id=self.parser_id,
                pages=[],
                text_extractable=True,
                error="python-docx is not installed",
            )

        try:
            document = Document(io.BytesIO(content))
            lines: list[str] = []

            for paragraph in document.paragraphs:
                text = " ".join(paragraph.text.split()).strip()
                if text:
                    lines.append(text)

            for table in document.tables:
                for row in table.rows:
                    cell_values = [" ".join(cell.text.split()).strip() for cell in row.cells]
                    row_text = " | ".join([value for value in cell_values if value])
                    if row_text:
                        lines.append(row_text)

            joined = "\n".join(lines).strip()
            pages = [ParsedPage(page=1, text=joined)] if joined else []

            return ParseResult(
                parser_id=self.parser_id,
                pages=pages,
                text_extractable=True,
            )
        except Exception as exc:
            return ParseResult(
                parser_id=self.parser_id,
                pages=[],
                text_extractable=True,
                error=f"docx parse failed: {exc}",
            )
